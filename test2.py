#coding=utf-8
__author__ = 'nanganglei'
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import tkMessageBox
import qrcode
import tkFileDialog
from Tkinter import *
root = Tk()
root.title("无线测试组--词表")
root.geometry("500x100")

def callback():
    global url_file_name
    entry.delete(0,END) #清空entry里面的内容
    #调用filedialog模块的askdirectory()函数去打开文件夹
    url_file_name =  tkFileDialog.askopenfilename()
    if url_file_name:
        entry.insert(0,url_file_name) #将选择好的路径加入到entry里面

def erweima():
    #打开文档
    document = Document()
    print url_file_name
    #加入不同等级的标题
    document.add_heading(u'生成二维码并记录对应URL的bug',0)
    #添加文本
    paragraph = document.add_paragraph(u'生成二维码！')
    # paragraph = document.add_paragraph('\n')
    urls_file = open(url_file_name)
    urls = urls_file.readlines()
    num = len(urls)
    index = 1
    for url in urls:
        paragraph = document.add_paragraph(str(index)+". "+url.strip())
        # paragraph = document.add_paragraph("\n")
        img = qrcode.make(url)
        img.save("erweima.png")
        document.add_picture('erweima.png', width=Inches(2))
        index = index + 1
    #设置字号
    # run = paragraph.add_run(u'设置字号、')
    # run.font.size = Pt(40)
    #设置字体
    # run = paragraph.add_run('Set Font,')
    # run.font.name = 'Consolas'
    #设置中文字体
    # run = paragraph.add_run(u'设置中文字体、')
    # run.font.name=u'宋体'
    # r = run._element
    # r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    #设置斜体
    # run = paragraph.add_run(u'斜体、')
    # run.italic = True
    #设置粗体
    # run = paragraph.add_run(u'粗体').bold = True
    #增加引用
    # document.add_paragraph('Intense quote', style='Intense Quote')
    #增加无序列表
    # document.add_paragraph(
    #     u'无序列表元素1', style='List Bullet'
    # )
    # document.add_paragraph(
    #     u'无序列表元素2', style='List Bullet'
    # )
    #增加有序列表
    # document.add_paragraph(
    #     u'有序列表元素1', style='List Number'
    # )
    # document.add_paragraph(
    #     u'有序列表元素2', style='List Number'
    # )
    #增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
    # document.add_picture('test.png', width=Inches(1.25))
    #增加表格
    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Name'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    #再增加3行表格元素
    # for i in xrange(3):
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = 'test'+str(i)
    #     row_cells[1].text = str(i)
    #     row_cells[2].text = 'desc'+str(i)
    #增加分页
    # document.add_page_break()
    #保存文件
    document.save(url_file_name.split("/")[-1] + '.docx')
    tkMessageBox.showinfo( 'tip','over!')


lable_select= Label(root,text="请选择包含url的文件：")
lable_select.grid(column=0, row=1,sticky=W)
entry = Entry(root,width=40)
entry.grid(column=0, row=2,sticky=W)

button = Button(root,text="Open",command=callback)
button.grid(column=1, row=2,sticky=W)

button_sure = Button(root,text="ok",command=erweima)
button_sure.grid(column=0, row=3,sticky=W)




root.mainloop()