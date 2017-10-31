#coding=utf-8
__author__ = 'nanganglei'
from docx import Document
from docx.shared import Inches
import tkMessageBox
import qrcode
import tkFileDialog
from Tkinter import *
import os
root = Tk()
root.title("二维码生成")
root.geometry("500x100")
rootdir = 'E:\\templates'
# rootdir = os.getcwd()
def callback():
    global url_file_name
    entry.delete(0,END) #清空entry里面的内容
    #调用filedialog模块的askdirectory()函数去打开文件夹
    url_file_name =  tkFileDialog.askopenfilename()
    if url_file_name:
        entry.insert(0,url_file_name) #将选择好的路径加入到entry里面

def erweima():
    #打开文档
    document = Document(os.path.join(rootdir, "default.docx"))
    # print url_file_name
    document.add_heading(u'批量生成二维码方便记录bug信息。',0)
    #添加文本
    paragraph = document.add_paragraph(u'生成二维码！-searchQA')
    # paragraph = document.add_paragraph('\n')
    urls_file = open(url_file_name)
    urls = urls_file.readlines()
    num = len(urls)
    index = 1
    for url in urls:
        paragraph = document.add_paragraph(str(index)+". "+url.strip())
        img = qrcode.make(url)
        img.save("erweima.png")
        document.add_picture('erweima.png', width=Inches(2))
        index = index + 1
    document.save(url_file_name.split("/")[-1] + '.docx')
    tkMessageBox.showinfo( 'tip','over!')
lable_select= Label(root,text="请选择包含url的文件：")
lable_select.grid(column=0, row=1,sticky=W)
entry = Entry(root,width=40)
entry.grid(column=0, row=2,sticky=W)

button = Button(root,text="Open",command=callback)
button.grid(column=1, row=2,sticky=W)

button_sure = Button(root,text="ok",command=erweima)
button_sure.grid(column=0, row=3,sticky=W)

root.mainloop()